# -*- coding: utf-8 -*-
"""
北京建筑大学新闻网 - 头条新闻爬虫
目标站点: https://xww.bucea.edu.cn/ttxw/index.htm
使用 DrissionPage + XPath 爬取，输出为 .docx 文件（每题以新闻标题命名）。
"""

import os
import re
import time
import sys
import random
from DrissionPage import SessionPage
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from urllib.parse import urljoin

# Windows 终端 UTF-8 支持
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ============================================================
# 配置
# ============================================================
BASE_URL = "https://xww.bucea.edu.cn"
LIST_URL = f"{BASE_URL}/ttxw/index.htm"
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")


def sanitize_filename(title: str) -> str:
    """去除 Windows 文件名中的非法字符，并限制长度。"""
    title = re.sub(r'[\\/*?:"<>|\n\r]', "", title)
    title = re.sub(r'\s+', " ", title).strip()
    if len(title) > 100:
        title = title[:100]
    return title


def get_list_page(page: SessionPage, url: str):
    """
    解析列表页，返回 (文章列表, 下一页URL)。
    文章列表每项: {title, url, date}
    """
    print(f"[列表页] 正在请求: {url}")

    for attempt in range(3):
        try:
            page.get(url)
        except Exception as e:
            print(f"[错误] 请求列表页失败: {e}，重试中... (attempt {attempt + 1})")
            time.sleep(2 * (attempt + 1))
            continue

        if page.html:
            break
        print(f"[警告] 列表页响应为空，重试中... (attempt {attempt + 1})")
        time.sleep(2 * (attempt + 1))
    else:
        print(f"[错误] 列表页请求失败，已重试3次，跳过此页")
        return [], None

    tree = etree.HTML(page.html)

    # XPath: 所有新闻条目
    media_items = tree.xpath('//div[@class="media"]')
    articles = []

    for item in media_items:
        # 标题 & 链接
        link_nodes = item.xpath('.//h4[@class="media-heading"]/a')
        if not link_nodes:
            continue

        href = link_nodes[0].get("href", "")
        # 优先使用 title 属性（完整标题），fallback 取文本
        title = (link_nodes[0].get("title", "") or "").strip()
        if not title:
            title = "".join(link_nodes[0].xpath(".//text()")).strip()

        full_url = urljoin(BASE_URL, href)

        # 日期
        date_nodes = item.xpath('.//div[@class="u-date"]/text()')
        date = ""
        for dn in date_nodes:
            d = dn.strip()
            if re.match(r'\d{4}-\d{2}-\d{2}', d):
                date = d
                break

        if title and href:
            articles.append({"title": title, "url": full_url, "date": date})

    # XPath: 下一页
    next_nodes = tree.xpath('//a[@class="Next"]/@href')
    next_url = None
    if next_nodes and next_nodes[0]:
        candidate = urljoin(url, next_nodes[0])
        # 避免死循环：如果下一页就是当前页，则停止
        if candidate != url:
            next_url = candidate

    return articles, next_url


def scrape_article(page: SessionPage, url: str, referer: str = ""):
    """
    抓取单篇新闻，返回 (标题, 正文文本, 元信息)。
    """
    original_referer = page.session.headers.get("Referer", "")
    if referer:
        page.session.headers["Referer"] = referer

    for attempt in range(3):
        try:
            page.get(url)
        except Exception as e:
            print(f"  [错误] 请求文章页失败: {e}")
            if attempt < 2:
                delay = 2 * (attempt + 1)
                print(f"  [重试] 等待 {delay}s 后重试... (attempt {attempt + 1})")
                time.sleep(delay)
            continue

        if page.html:
            break
        print(f"  [警告] 页面响应为空，重试中... (attempt {attempt + 1})")
        if attempt < 2:
            time.sleep(2 * (attempt + 1))
    else:
        print(f"  [警告] 重试后页面仍为空，跳过")
        page.session.headers["Referer"] = original_referer
        return "", "", ""

    page.session.headers["Referer"] = original_referer

    tree = etree.HTML(page.html)

    # 标题：<h2 class="detail-title">
    title_nodes = tree.xpath('//h2[@class="detail-title"]/text()')
    title = "".join(title_nodes).strip() if title_nodes else "未知标题"

    # 正文：<div class="m-content" id="art">
    content_parts = []
    content_divs = tree.xpath('//div[@class="m-content"]')
    if content_divs:
        for p in content_divs[0].xpath('.//p'):
            text = "".join(p.xpath(".//text()")).strip()
            if text:
                content_parts.append(text)

    content = "\n".join(content_parts)

    # 元信息
    meta_nodes = tree.xpath('//p/small')
    if meta_nodes:
        meta_text = "".join(meta_nodes[0].xpath(".//text()")).strip()
        meta_text = re.sub(r'\s+', " ", meta_text)
    else:
        meta_text = ""

    return title, content, meta_text


def save_as_docx(title: str, content: str, meta: str, date: str) -> str:
    """将新闻保存为 .docx 文件，返回保存路径。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    # 大标题
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.size = Pt(18)

    # 元信息
    if meta or date:
        meta_line = f"发布时间: {date}" if date else meta
        p = doc.add_paragraph()
        run = p.add_run(meta_line)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # 正文
    for para in content.split("\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(14)

    filename = sanitize_filename(title) + ".docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    # 处理重名文件
    counter = 1
    while os.path.exists(filepath):
        filepath = os.path.join(OUTPUT_DIR, f"{sanitize_filename(title)}_{counter}.docx")
        counter += 1

    doc.save(filepath)
    return filepath


def load_scraped_titles(output_dir: str) -> set:
    """扫描 output 目录中已有的 .docx 文件，返回已爬取标题的集合（sanitize 后的文件名）。"""
    scraped = set()
    if not os.path.exists(output_dir):
        return scraped
    for filename in os.listdir(output_dir):
        if not filename.endswith(".docx"):
            continue
        name = filename[:-5]  # 去掉 .docx
        # 去掉防重名后缀 _1, _2 等，还原为 sanitize 后的基础名
        base = re.sub(r'_\d+$', '', name)
        scraped.add(base)
    return scraped


# ============================================================
# 主流程
# ============================================================
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    page = SessionPage()
    page.session.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/125.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9",
    })

    # 可配置：从命令行参数控制最大页数
    # 用法: python scraper.py [max_pages]
    # 例如: python scraper.py 3  只爬3页
    #       python scraper.py    爬全部65页
    max_pages = int(sys.argv[1]) if len(sys.argv) > 1 else 0  # 0 表示不限制

    scraped_titles = load_scraped_titles(OUTPUT_DIR)

    current_page_url = LIST_URL
    page_num = 1
    total_saved = 0
    total_skipped = 0

    print("=" * 60)
    print("  北京建筑大学新闻网 - 头条新闻爬虫")
    print(f"  目标: {LIST_URL}")
    print(f"  输出目录: {OUTPUT_DIR}")
    print(f"  已爬取: {len(scraped_titles)} 篇")
    if max_pages:
        print(f"  限制页数: {max_pages} 页")
    print("=" * 60)

    while current_page_url:
        print(f"\n{'=' * 50}")
        print(f">> 第 {page_num} 页")

        articles, next_url = get_list_page(page, current_page_url)

        if not articles:
            print("  未找到任何文章，退出循环。")
            break

        # 过滤已爬取的文章
        new_articles = []
        for a in articles:
            if sanitize_filename(a["title"]) not in scraped_titles:
                new_articles.append(a)

        skipped = len(articles) - len(new_articles)
        total_skipped += skipped
        print(f"  本页共 {len(articles)} 篇（已爬 {skipped} 篇，待爬 {len(new_articles)} 篇）")

        for idx, article in enumerate(new_articles, 1):
            title_preview = (
                article["title"][:60] + "..."
                if len(article["title"]) > 60
                else article["title"]
            )
            print(f"  [{idx:2d}/{len(new_articles)}] 正在爬取: {title_preview}")

            try:
                title, content, meta = scrape_article(page, article["url"], current_page_url)

                if not content:
                    print(f"        [警告] 正文为空，跳过")
                    continue

                saved_path = save_as_docx(title, content, meta, article["date"])
                total_saved += 1
                scraped_titles.add(sanitize_filename(title))
                print(f"        [已保存] {os.path.basename(saved_path)}")

            except Exception as e:
                print(f"        [异常] {e}")
                continue

            time.sleep(1.5 + random.random() * 1.5)  # 1.5-3s 随机间隔

        # 检查是否继续翻页
        page_num += 1
        if max_pages and page_num > max_pages:
            print(f"\n  已达到最大页数限制 ({max_pages} 页)，停止翻页。")
            break

        if next_url:
            print(f"\n  --> 发现下一页: {next_url}")
            current_page_url = next_url
            time.sleep(1.5 + random.random() * 1.5)
        else:
            print(f"\n  --> 已到达最后一页")
            break

    print(f"\n{'=' * 60}")
    print(f"  爬取完成! 本批保存 {total_saved} 篇，跳过 {total_skipped} 篇已爬")
    print(f"  输出目录: {OUTPUT_DIR}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
